import streamlit as st
import pickle
import re
import nltk
import pdfplumber
import os
from docx import Document
from nltk.corpus import stopwords
from suggest import role_profiles

nltk.download('stopwords', quiet=True)
stop_words = set(stopwords.words('english'))

@st.cache_resource
def load_resources():
    try:
        vectorizer = pickle.load(open("vectorizer.pkl", "rb"))
        model = pickle.load(open("model.pkl", "rb"))
        encoder = pickle.load(open("encoder.pkl", "rb"))
        return vectorizer, model, encoder
    except Exception as e:
        st.error(f"Error loading resources: {e}")
        return None, None, None

def extract_text(file, file_type):
    try:
        if file_type == 'pdf':
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text
        elif file_type == 'docx':
            doc = Document(file)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_type == 'txt':
            return file.read().decode('utf-8', errors='ignore')
        return ''
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return ''

def clean_resume(txt):
    txt = re.sub(r"http\S+|www\S+|https\S+", '', txt)
    txt = re.sub(r"\S+@\S+", '', txt)
    txt = re.sub(r"<.*?>", '', txt)
    txt = re.sub(r"[\r\n]+", ' ', txt)
    txt = txt.lower()
    txt = ' '.join([word for word in txt.split() if word not in stop_words])
    return txt

# =======================
# Main App
# =======================
st.title("📄 Resume Screening & Suggestions")

upload_file = st.file_uploader("📤 Upload Resume", type=["pdf", "docx", "txt"])
available_categories = list(role_profiles.keys())
selected_category = st.selectbox("🎯 Select Target Job Category", available_categories)

if upload_file and selected_category:
    file_type = upload_file.name.split('.')[-1]
    resume_text = extract_text(upload_file, file_type)

    if not resume_text.strip():
        st.warning("⚠️ Could not extract readable text.")
        st.stop()

    cleaned_resume = clean_resume(resume_text)

    vectorizer, model, encoder = load_resources()
    if not all([vectorizer, model, encoder]):
        st.stop()

    try:
        vectorized_resume = vectorizer.encode([cleaned_resume])
    except Exception as e:
        st.error(f"Error during vectorization: {e}")
        st.stop()

    try:
        if hasattr(model, "predict_proba"):
            probas = model.predict_proba(vectorized_resume)[0]
            predicted_idx = probas.argmax()
            predicted_category = encoder.inverse_transform([predicted_idx])[0]
            confidence = probas[predicted_idx] * 100
            st.info(f"🤖 Model Prediction: **{predicted_category}** ({confidence:.2f}% confidence)")
        else:
            predicted_idx = model.predict(vectorized_resume)[0]
            predicted_category = encoder.inverse_transform([predicted_idx])[0]
            st.info(f"🤖 Model Prediction: **{predicted_category}**")
    except Exception as e:
        st.error(f"Error during prediction: {e}")
        st.stop()

    if predicted_category != selected_category:
        st.warning(f"⚠️ Your resume seems more aligned with **{predicted_category}** than **{selected_category}**.")

    suggestions = role_profiles.get(selected_category, {}).get("tips", [])
    if suggestions:
        st.markdown(f"### ✅ Suggestions to Improve for **{selected_category}** Role:")
        for tip in suggestions:
            st.markdown(f"- {tip}")
    else:
        st.warning("⚠️ No suggestions available.")

    # Store resume for editor
    st.session_state["resume_text"] = resume_text

    # =======================
    # Edit Resume with Auto-Create Page
    # =======================
    if st.button("✏️ Edit Resume"):
        pages_dir = os.path.join(os.path.dirname(__file__), "pages")
        editor_page_path = os.path.join(pages_dir, "editor_page.py")

        # Create pages folder if missing
        if not os.path.exists(pages_dir):
            os.makedirs(pages_dir)

        # Create editor_page.py if missing
        if not os.path.exists(editor_page_path):
            default_editor_code = '''import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches
from editor import resume_editor

st.set_page_config(page_title="Resume Editor", page_icon="✏️")

st.title("✏️ Resume Editor")

if "resume_text" not in st.session_state:
    st.error("⚠️ No resume loaded. Please go back to the main page.")
    if st.button("⬅️ Back to Suggestions"):
        st.switch_page("app.py")
    st.stop()

resume_text = st.session_state["resume_text"]

edited_resume_text, uploaded_image = resume_editor(resume_text)

st.subheader("📥 Download Updated Resume")
col1, col2 = st.columns(2)

with col1:
    txt_bytes = edited_resume_text.encode("utf-8")
    st.download_button(
        label="Download as TXT",
        data=txt_bytes,
        file_name="Edited_Resume.txt",
        mime="text/plain"
    )

with col2:
    doc_buffer = BytesIO()
    doc = Document()
    if uploaded_image:
        doc.add_picture(uploaded_image, width=Inches(1.5))
    doc.add_paragraph(edited_resume_text)
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    st.download_button(
        label="Download as DOCX",
        data=doc_buffer,
        file_name="Edited_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if st.button("⬅️ Back to Suggestions"):
    st.switch_page("app.py")
'''
            with open(editor_page_path, "w", encoding="utf-8") as f:
                f.write(default_editor_code)

        # Now switch to editor page
        st.switch_page("pages/editor_page.py")
