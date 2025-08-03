import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches
from editor import resume_editor  # ✅ Import engaging editor

st.set_page_config(page_title="Resume Editor", page_icon="✏️")

st.title("✏️ Resume Editor")

# =======================
# Check if resume exists
# =======================
if "resume_text" not in st.session_state:
    st.error("⚠️ No resume loaded. Please go back to the main page.")
    if st.button("⬅️ Back to Suggestions"):
        st.switch_page("app.py")
    st.stop()

# Load resume from session
resume_text = st.session_state["resume_text"]

# =======================
# Resume Editor
# =======================
edited_resume_text, uploaded_image = resume_editor(resume_text)

# =======================
# Download Updated Resume
# =======================
st.subheader("📥 Download Updated Resume")
col1, col2 = st.columns(2)

with col1:
    txt_bytes = edited_resume_text.encode("utf-8")
    st.download_button(
        label="Download as TXT",
        data=txt_bytes,
        file_name="Edited_Resume.txt",
        mime="text/plain"
    )

with col2:
    doc_buffer = BytesIO()
    doc = Document()
    if uploaded_image:
        doc.add_picture(uploaded_image, width=Inches(1.5))
    doc.add_paragraph(edited_resume_text)
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    st.download_button(
        label="Download as DOCX",
        data=doc_buffer,
        file_name="Edited_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =======================
# Back Button
# =======================
if st.button("⬅️ Back to Suggestions"):
    st.switch_page("app.py")
