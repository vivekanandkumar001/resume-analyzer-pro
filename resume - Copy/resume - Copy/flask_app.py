
from flask import Flask, request, jsonify, send_file
import io
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from resume_templates import generate_resume_pdf
from model import predict_category_and_conf
from suggest import analyze_for_role, suggest_from_resume, log_feedback_rows
import pdfplumber
import docx

app = Flask(__name__)

# ----------- Helpers -----------
def _extract_text_from_upload(file_storage):
    """
    Robust extractor for Flask uploads (PDF/DOCX/TXT).
    Does not depend on parser.py's .name attribute.
    """
    filename = (file_storage.filename or "").lower()
    raw = file_storage.read()

    if filename.endswith(".pdf"):
        text = ""
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                text += t + "\n"
        return text.strip()

    if filename.endswith(".docx"):
        d = docx.Document(io.BytesIO(raw))
        return "\n".join(p.text for p in d.paragraphs).strip()

    # fallback to txt
    try:
        return raw.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def _make_pdf(text: str) -> io.BytesIO:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica", 11)
    for line in (text or "").splitlines():
        # naive wrapping at ~95 chars
        while len(line) > 95:
            c.drawString(50, y, line[:95])
            line = line[95:]
            y -= 14
            if y < 60:
                c.showPage(); c.setFont("Helvetica", 11); y = height - 50
        c.drawString(50, y, line)
        y -= 14
        if y < 60:
            c.showPage(); c.setFont("Helvetica", 11); y = height - 50
    c.showPage()
    c.save()
    buf.seek(0)
    return buf


def _make_docx(text: str) -> io.BytesIO:
    buf = io.BytesIO()
    d = Document()
    for para in (text or "").splitlines():
        d.add_paragraph(para)
    d.save(buf)
    buf.seek(0)
    return buf


@app.route("/")
def home():
    return jsonify({"message": "AI Resume Analyzer Flask API is running."})


# 1) Upload + Predict
@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    text = _extract_text_from_upload(file)
    if not text:
        return jsonify({"error": "Could not extract text from file"}), 400
    category, conf = predict_category_and_conf(text)
    return jsonify({
        "resume_text": text,
        "predicted_category": category,
        "confidence": conf
    })


# 2) Analyze against a chosen category
@app.route("/analyze", methods=["POST"])
def analyze_resume():
    data = request.json or {}
    text = data.get("resume_text")
    role = data.get("category")
    if not text or not role:
        return jsonify({"error": "resume_text and category required"}), 400
    result = analyze_for_role(text, role)
    return jsonify(result)


# 3) Editor-style re-analyze (same as analyze but named for clarity)
@app.route("/editor/analyze", methods=["POST"])
def editor_analyze():
    data = request.json or {}
    text = data.get("resume_text")
    role = data.get("category")
    if not text or not role:
        return jsonify({"error": "resume_text and category required"}), 400
    result = suggest_from_resume(text, role)
    return jsonify(result)


# 4) Re-check (predict category + analysis)
@app.route("/editor/recheck", methods=["POST"])
def recheck_resume():
    data = request.json or {}
    text = data.get("resume_text")
    role = data.get("category")
    if not text or not role:
        return jsonify({"error": "resume_text and category required"}), 400
    pred_cat, conf = predict_category_and_conf(text)
    res = analyze_for_role(text, role)
    return jsonify({
        "model_thinks": pred_cat,
        "confidence": conf,
        "analysis": res
    })


# 5) Download via simple PDF/DOCX/TXT
@app.route("/editor/download", methods=["POST"])
def download_resume_editor():
    data = request.json or {}
    text = data.get("resume_text", "")
    fmt = (data.get("format") or "pdf").lower()

    if fmt == "pdf":
        buf = _make_pdf(text)
        return send_file(buf, as_attachment=True,
                         download_name="Updated_Resume.pdf",
                         mimetype="application/pdf")
    if fmt == "docx":
        buf = _make_docx(text)
        return send_file(buf, as_attachment=True,
                         download_name="Updated_Resume.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if fmt == "txt":
        return send_file(io.BytesIO(text.encode("utf-8")),
                         as_attachment=True,
                         download_name="Updated_Resume.txt",
                         mimetype="text/plain")
    return jsonify({"error": "Unsupported format"}), 400


# 6) Template-based PDF (uses reportlab template variants)
@app.route("/download", methods=["POST"])
def download_template_pdf():
    data = request.json or {}
    text = data.get("resume_text", "")
    try:
        template_id = int(data.get("template_id", 1))
    except Exception:
        template_id = 1

    out = io.BytesIO()
    generate_resume_pdf(text, template_id, out)
    out.seek(0)
    return send_file(out, as_attachment=True,
                     download_name="resume.pdf",
                     mimetype="application/pdf")


# 7) Feedback logging + lightweight RL tick
@app.route("/feedback", methods=["POST"])
def feedback():
    data = request.json or {}
    text = data.get("resume_text", "")
    role = data.get("category", "")
    reward = int(data.get("reward", 1))  # 1 or -1
    comments = data.get("comments", "")

    # Basic safety
    if not role:
        return jsonify({"error": "category required"}), 400

    # Build rows from current analysis
    analysis = analyze_for_role(text, role)
    rows = []
    for s in analysis.get("missing_skills", []):
        rows.append(("skill", s))
    for p in analysis.get("projects", []):
        rows.append(("project", p))
    for c in analysis.get("courses", []):
        rows.append(("course", c))
    for c in analysis.get("certificates", []):
        rows.append(("certificate", c))

    log_feedback_rows(
        resume_id="session", target_role=role, rows=rows,
        reward=reward, comments=comments
    )

    # live-tick RL (optional)
    try:
        from train_rl_from_feedback import train_incremental
        train_incremental()
        return jsonify({"status": "ok", "note": "feedback logged & RL updated"})
    except Exception as e:
        return jsonify({"status": "ok", "note": f"feedback logged; RL update skipped: {e}"}), 200


if __name__ == "__main__":
    # Use FLASK_RUN_PORT/FLASK_RUN_HOST when running via `flask run`
    app.run(host="0.0.0.0", port=5000, debug=True)
